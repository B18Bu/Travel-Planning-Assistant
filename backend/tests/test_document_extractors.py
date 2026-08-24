import base64
from uuid import uuid4

import pytest

try:
    from docx import Document
except ImportError:
    Document = None
try:
    import fitz
except ImportError:
    fitz = None

from app.services.document_extractors import (
    chunk_extracted_content,
    extract_docx,
    extract_pdf_with_pymupdf,
)


@pytest.mark.skipif(Document is None, reason="未安装 python-docx")
def test_extract_docx_keeps_heading_paths_body_order_and_semantic_table(tmp_path):
    path = tmp_path / "report.docx"
    document = Document()
    document.add_paragraph("前言")
    document.add_heading("行程方案", level=1)
    document.add_paragraph("概览")
    document.add_heading("第一天", level=2)
    document.add_paragraph("上午游览")
    table = document.add_table(rows=3, cols=3)
    for row, values in zip(table.rows, (("景点", "时段", ""), ("古城", "上午", ""), ("博物馆", "", "室内"))):
        for cell, value in zip(row.cells, values):
            cell.text = value
    document.add_heading("注意事项", level=1)
    document.add_paragraph("携带雨具")
    document.save(path)

    extracted = extract_docx(path, tmp_path / "extracted")

    assert [item["content"] for item in extracted] == [
        "前言", "行程方案", "概览", "第一天", "上午游览",
        "章节：行程方案 > 第一天\n表格 1\n表头：景点 | 时段 | 第 3 列\n第 1 行：景点=古城；时段=上午；第 3 列=\n第 2 行：景点=博物馆；时段=；第 3 列=室内",
        "注意事项", "携带雨具",
    ]
    assert [item["section_path"] for item in extracted] == [
        (), ("行程方案",), ("行程方案",), ("行程方案", "第一天"),
        ("行程方案", "第一天"), ("行程方案", "第一天"),
        ("注意事项",), ("注意事项",),
    ]
    assert [item["source_section"] for item in extracted] == [
        "正文", "行程方案", "行程方案", "行程方案 > 第一天",
        "行程方案 > 第一天", "行程方案 > 第一天", "注意事项", "注意事项",
    ]
    assert [item["source_order"] for item in extracted] == list(range(8))
    assert extracted[5]["source_table"] == 1


def test_chunking_aggregates_adjacent_text_but_keeps_tables_as_boundaries():
    document_id = uuid4()
    items = [
        {"content": "甲段", "chunk_type": "text", "section_path": (), "source_section": "正文", "source_order": 1},
        {"content": "乙段", "chunk_type": "text", "section_path": (), "source_section": "正文", "source_order": 2},
        {"content": "章节：正文\n表格 1\n表头：项目\n第 2 行：项目=值", "chunk_type": "table", "source_section": "正文", "source_table": 1, "source_order": 3},
        {"content": "丙段", "chunk_type": "text", "section_path": (), "source_section": "正文", "source_order": 4},
    ]

    chunks = chunk_extracted_content(document_id, "报告.docx", items)

    assert [chunk.content for chunk in chunks] == [
        "甲段\n\n乙段", "章节：正文\n表格 1\n表头：项目\n第 2 行：项目=值", "丙段",
    ]
    assert [chunk.chunk_type for chunk in chunks] == ["text", "table", "text"]
    assert [(chunk.char_start, chunk.char_end) for chunk in chunks] == [(0, 6), (0, 10), (0, 2)]


def test_chunking_prefers_sentence_boundaries_and_tracks_body_offsets():
    document_id = uuid4()
    body = "甲" * 12 + "。" + "乙" * 12 + "。"

    chunks = chunk_extracted_content(
        document_id, "报告.docx",
        [{"content": body, "chunk_type": "text", "section_path": ("行程",), "source_section": "行程", "source_order": 1}],
        max_chars=30, overlap=5,
    )

    assert [chunk.content for chunk in chunks] == ["章节：行程\n\n甲甲甲甲甲甲甲甲甲甲甲甲。", "章节：行程\n\n甲甲甲甲。乙乙乙乙乙乙乙乙乙乙乙乙。"]
    assert [(chunk.char_start, chunk.char_end) for chunk in chunks] == [(0, 13), (8, 26)]
    assert all(len(chunk.content) <= 30 for chunk in chunks)


def test_chunking_splits_tables_by_rows_with_repeated_context_and_continues_long_rows():
    document_id = uuid4()
    content = (
        "章节：行程\n表格 2\n表头：景点 | 说明\n"
        "第 2 行：景点=古城；说明=推荐。\n"
        f"第 3 行：景点=博物馆；说明={'甲' * 50}"
    )

    chunks = chunk_extracted_content(
        document_id, "报告.docx",
        [{"content": content, "chunk_type": "table", "source_section": "行程", "source_table": 2, "source_order": 1}],
        max_chars=70, overlap=5,
    )

    assert len(chunks) >= 3
    assert all(chunk.content.startswith("章节：行程\n表格 2\n表头：景点 | 说明\n") for chunk in chunks)
    assert chunks[0].content.endswith("第 2 行：景点=古城；说明=推荐。")
    assert any("第 3 行（续）" in chunk.content for chunk in chunks[1:])
    assert all(chunk.chunk_type == "table" and len(chunk.content) <= 70 for chunk in chunks)


def test_chunking_chart_ocr_gets_its_own_context_and_text_overlap():
    document_id = uuid4()
    ocr_text = "图" * 30

    chunks = chunk_extracted_content(
        document_id, "报告.docx",
        [
            {"content": "前文", "chunk_type": "text", "section_path": ("行程",), "source_section": "行程", "source_order": 1},
            {"content": ocr_text, "chunk_type": "chart_ocr", "source_section": "行程", "source_figure": 4, "source_order": 2},
        ],
        max_chars=30, overlap=5,
    )

    assert chunks[0].content == "章节：行程\n\n前文"
    chart_chunks = chunks[1:]
    assert all(chunk.chunk_type == "chart_ocr" for chunk in chart_chunks)
    assert all(chunk.content.startswith("章节：行程\n图表 4\n\n") for chunk in chart_chunks)
    assert chart_chunks[0].content[-5:] == chart_chunks[1].content[len("章节：行程\n图表 4\n\n"):][:5]
    assert [(chunk.source_section, chunk.source_figure) for chunk in chart_chunks] == [("行程", 4)] * len(chart_chunks)


def test_chunking_rejects_a_limit_or_overlap_larger_than_contract():
    with pytest.raises(ValueError, match="不得超过 800"):
        chunk_extracted_content(
            document_id=uuid4(), document_name="报告.docx",
            items=[{"content": "正文", "chunk_type": "text"}], max_chars=801,
        )
    with pytest.raises(ValueError, match="不得超过 100"):
        chunk_extracted_content(
            document_id=uuid4(), document_name="报告.docx",
            items=[{"content": "正文", "chunk_type": "text"}], overlap=101,
        )


@pytest.mark.parametrize(
    ("content", "max_chars", "expected_first"),
    [
        ("甲\n\n乙丙", 4, "甲\n\n"),
        ("甲！乙丙", 3, "甲！"),
        ("甲？乙丙", 3, "甲？"),
        ("甲；乙丙", 3, "甲；"),
        ("甲\n乙丙", 3, "甲\n"),
        ("甲，乙丙", 3, "甲，乙"),
    ],
)
def test_chunking_uses_required_natural_boundaries_without_comma(content, max_chars, expected_first):
    chunks = chunk_extracted_content(
        uuid4(), "报告.docx", [{"content": content, "chunk_type": "text"}],
        max_chars=max_chars, overlap=0,
    )

    assert chunks[0].content == expected_first


def test_chunking_uses_default_800_limit_and_advances_after_short_paragraph_boundary():
    long_chunks = chunk_extracted_content(
        uuid4(), "报告.docx", [{"content": "甲" * 900, "chunk_type": "text"}],
    )
    assert [len(chunk.content) for chunk in long_chunks] == [800, 200]

    import threading

    result: list[object] = []
    worker = threading.Thread(
        target=lambda: result.extend(chunk_extracted_content(
            uuid4(), "报告.docx", [{"content": "甲\n\n乙丙", "chunk_type": "text"}],
            max_chars=4, overlap=3,
        )),
        daemon=True,
    )
    worker.start()
    worker.join(timeout=0.2)

    assert not worker.is_alive(), "短段落切点不得导致分块循环停滞"
    assert [chunk.content for chunk in result] == ["甲\n\n", "乙丙"]
    assert [(chunk.char_start, chunk.char_end) for chunk in result] == [(0, 3), (3, 5)]


def test_chunking_prefixes_non_body_text_without_source_order():
    chunks = chunk_extracted_content(
        uuid4(), "报告.docx",
        [{"content": "内容", "chunk_type": "text", "section_path": ("行程",), "source_section": "行程"}],
    )

    assert [chunk.content for chunk in chunks] == ["章节：行程\n\n内容"]


def test_table_chunks_repeat_context_without_overlapping_complete_data_rows():
    context = "章节：行程\n表格 1\n表头：项目"
    rows = [f"第 {index} 行：项目=数据{index}" for index in range(2, 6)]
    chunks = chunk_extracted_content(
        uuid4(), "报告.docx",
        [{"content": "\n".join([context, *rows]), "chunk_type": "table", "source_section": "行程", "source_table": 1}],
        max_chars=35, overlap=5,
    )

    assert len(chunks) == 4
    assert all(chunk.content.startswith(f"{context}\n") for chunk in chunks)
    combined = "\n".join(chunk.content.removeprefix(f"{context}\n") for chunk in chunks)
    assert all(combined.count(row) == 1 for row in rows)


def test_chunking_truncates_oversized_prefixes_but_keeps_last_section_title():
    section = "前级" * 40 + " > 末级标题"
    text_chunks = chunk_extracted_content(
        uuid4(), "报告.docx",
        [{"content": "正文", "chunk_type": "text", "source_section": section}],
        max_chars=20, overlap=0,
    )
    chart_chunks = chunk_extracted_content(
        uuid4(), "报告.docx",
        [{"content": "图表正文", "chunk_type": "chart_ocr", "source_section": section, "source_figure": 1}],
        max_chars=20, overlap=0,
    )
    table_chunks = chunk_extracted_content(
        uuid4(), "报告.docx",
        [{"content": f"章节：{section}\n表格 1\n表头：列\n第 1 行：列=值", "chunk_type": "table", "source_section": section, "source_table": 1}],
        max_chars=20, overlap=0,
    )

    assert all(chunk.content and len(chunk.content) <= 20 for chunk in text_chunks + chart_chunks + table_chunks)
    assert all("末级标题" in chunk.content for chunk in text_chunks + chart_chunks + table_chunks)
    assert text_chunks[-1].content.endswith("正文")
    assert chart_chunks[-1].content.endswith("图表正文")
    assert table_chunks[-1].content.endswith("值")


def test_table_chunk_offsets_follow_original_row_body_and_ignore_continuation_labels():
    context = "章节：行程\n表格 1\n表头：项目"
    first = "第 1 行：项目=短值"
    second = "第 2 行：项目=" + "甲" * 30
    body = f"{first}\n{second}"
    chunks = chunk_extracted_content(
        uuid4(), "报告.docx",
        [{"content": f"{context}\n{body}", "chunk_type": "table", "source_section": "行程", "source_table": 1}],
        max_chars=42, overlap=0,
    )

    assert len(chunks) >= 3
    assert (chunks[0].char_start, chunks[0].char_end) == (0, len(first))
    second_start = len(first) + 1
    assert chunks[1].char_start == second_start + len("第 2 行：")
    assert chunks[-1].char_start < chunks[-1].char_end == second_start + len(second)
    assert all(second_start <= chunk.char_start < chunk.char_end <= second_start + len(second) for chunk in chunks[1:])


def test_chunking_keeps_compact_last_section_and_safe_source_section_for_tiny_limits():
    section = "前级" * 80 + " > 末级标题"
    items = [
        {"content": "正文", "chunk_type": "text", "source_section": section},
        {"content": "图正文", "chunk_type": "chart_ocr", "source_section": section, "source_figure": 1},
        {"content": f"章节：{section}\n表格 1\n表头：列\n第 1 行：列=值", "chunk_type": "table", "source_section": section, "source_table": 1},
    ]

    chunks = chunk_extracted_content(uuid4(), "报告.docx", items, max_chars=15, overlap=0)

    assert all(0 < len(chunk.content) <= 15 for chunk in chunks)
    assert all("末级标题" in chunk.content and "末级标题" in chunk.source_section for chunk in chunks)
    assert all(len(chunk.source_section) <= 100 for chunk in chunks)


def test_chunking_aggregates_adjacent_same_section_text_without_source_order():
    chunks = chunk_extracted_content(
        uuid4(), "报告.docx",
        [
            {"content": "甲", "chunk_type": "text", "section_path": ("行程",), "source_section": "行程"},
            {"content": "乙", "chunk_type": "text", "section_path": ("行程",), "source_section": "行程"},
            {"content": "丙", "chunk_type": "text", "section_path": ("行程",), "source_section": "行程", "source_order": 3},
            {"content": "丁", "chunk_type": "text", "section_path": ("行程",), "source_section": "行程", "source_order": 5},
        ],
    )

    assert [chunk.content for chunk in chunks] == ["章节：行程\n\n甲\n\n乙\n\n丙", "章节：行程\n\n丁"]


@pytest.mark.skipif(Document is None, reason="未安装 python-docx")
def test_extract_docx_flattens_cell_newlines_before_table_chunking(tmp_path):
    document = Document()
    table = document.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "项目"
    table.cell(1, 0).text = "古城\n夜游"
    path = tmp_path / "multiline-table.docx"
    document.save(path)

    extracted = extract_docx(path, tmp_path / "extracted")
    table_item = next(item for item in extracted if item["chunk_type"] == "table")
    chunks = chunk_extracted_content(uuid4(), "报告.docx", [table_item], max_chars=25, overlap=0)

    assert table_item["content"] == "章节：正文\n表格 1\n表头：项目\n第 1 行：项目=古城 夜游"
    assert all("第 1 行：" not in chunk.content or "\n夜游" not in chunk.content for chunk in chunks)
    assert all(len(chunk.content) <= 25 for chunk in chunks)


@pytest.mark.skipif(Document is None, reason="未安装 python-docx")
def test_extract_docx_truncates_same_heading_level_and_recognizes_chinese_heading_style(tmp_path):
    from docx.enum.style import WD_STYLE_TYPE

    document = Document()
    chinese_heading = document.styles.add_style("标题 2", WD_STYLE_TYPE.PARAGRAPH)
    chinese_heading.base_style = document.styles["Heading 2"]
    document.add_heading("一级甲", level=1)
    document.add_heading("二级甲", level=2)
    document.add_heading("二级乙", level=2)
    document.add_heading("一级乙", level=1)
    paragraph = document.add_paragraph("中文二级")
    paragraph.style = chinese_heading
    document.add_paragraph("正文")
    path = tmp_path / "headings.docx"
    document.save(path)

    extracted = extract_docx(path, tmp_path / "extracted")

    assert [item["section_path"] for item in extracted] == [
        ("一级甲",), ("一级甲", "二级甲"), ("一级甲", "二级乙"),
        ("一级乙",), ("一级乙", "中文二级"), ("一级乙", "中文二级"),
    ]


def test_chunking_keeps_distinct_source_sections_without_section_paths_separate():
    chunks = chunk_extracted_content(
        uuid4(), "报告.pdf",
        [
            {"content": "甲", "chunk_type": "text", "source_section": "甲章"},
            {"content": "乙", "chunk_type": "text", "source_section": "乙章"},
        ],
    )

    assert [chunk.content for chunk in chunks] == ["章节：甲章\n\n甲", "章节：乙章\n\n乙"]
    assert [chunk.source_section for chunk in chunks] == ["甲章", "乙章"]


@pytest.mark.skipif(fitz is None, reason="未安装 PyMuPDF")
def test_pdf_pages_remain_separate_during_chunking(tmp_path):
    pdf_path = tmp_path / "two-pages.pdf"
    document = fitz.open()
    page_one = document.new_page()
    page_one.insert_text((72, 72), "First page")
    page_two = document.new_page()
    page_two.insert_text((72, 72), "Second page")
    document.save(pdf_path)
    document.close()

    extracted = extract_pdf_with_pymupdf(pdf_path, tmp_path / "extracted")
    chunks = chunk_extracted_content(uuid4(), "two-pages.pdf", extracted)

    assert [chunk.content for chunk in chunks] == ["First page", "Second page"]
    assert [chunk.source_page for chunk in chunks] == [1, 2]
    assert [(chunk.char_start, chunk.char_end) for chunk in chunks] == [(0, 10), (0, 11)]


@pytest.mark.skipif(fitz is None, reason="未安装 PyMuPDF")
def test_extract_pdf_with_pymupdf_returns_page_text_and_exports_images(tmp_path):
    pdf_path = tmp_path / "report.pdf"
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Travel research")
    page.insert_image(fitz.Rect(72, 100, 100, 128), stream=base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAIAAACQd1PeAAAADElEQVR4nGP4z8AAAAMBAQDJ/pLvAAAAAElFTkSuQmCC"
    ))
    document.save(pdf_path)
    document.close()

    extracted = extract_pdf_with_pymupdf(pdf_path, tmp_path / "extracted")

    assert extracted[0] == {"content": "Travel research", "chunk_type": "text", "source_page": 1}
    image = extracted[1]
    assert image["chunk_type"] == "chart_ocr"
    assert image["source_page"] == 1
    assert image["source_figure"] == 1
    assert (tmp_path / "extracted" / image["image_path"]).exists()
