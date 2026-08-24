from __future__ import annotations

import re
from pathlib import Path
from uuid import UUID, uuid4

from app.models.documents import CHUNK_CONTENT_MAX_LENGTH, DocumentChunk

MAX_CHARS = CHUNK_CONTENT_MAX_LENGTH
OVERLAP = 100


def _section_name(path: tuple[str, ...]) -> str:
    return " > ".join(path) if path else "正文"


def _natural_end(content: str, start: int, limit: int) -> int:
    """返回不超过上限且优先保留结构边界的结束位置。"""
    maximum = min(start + limit, len(content))
    if maximum == len(content):
        return maximum
    window = content[start:maximum]
    paragraph_end = window.rfind("\n\n")
    if paragraph_end >= 0:
        return start + paragraph_end + 2
    sentence_end = max(window.rfind(character) for character in "。！？；")
    if sentence_end >= 0:
        return start + sentence_end + 1
    newline_end = window.rfind("\n")
    if newline_end >= 0:
        return start + newline_end + 1
    return maximum


def _compact_section(section: object, limit: int = 100) -> str | None:
    if not isinstance(section, str) or not section:
        return None
    if len(section) <= limit:
        return section
    last_section = section.rsplit(" > ", 1)[-1]
    if len(last_section) >= limit:
        return last_section[-limit:]
    return f"… > {last_section}"[-limit:]


def _make_chunk(
    document_id: UUID,
    document_name: str,
    content: str,
    item: dict,
    *,
    char_start: int,
    char_end: int,
) -> DocumentChunk:
    return DocumentChunk(
        id=str(uuid4()),
        document_id=str(document_id),
        content=content,
        chunk_type=item["chunk_type"],
        document_name=document_name,
        source_page=item.get("source_page"),
        source_section=_compact_section(item.get("source_section")),
        source_table=item.get("source_table"),
        source_figure=item.get("source_figure"),
        image_path=item.get("image_path"),
        char_start=char_start,
        char_end=char_end,
    )


def _prefix_with_section(section: str, suffix: str, max_chars: int) -> str:
    """压缩来源前缀，保留末级标题并为正文至少留一个字符。"""
    last_section = section.rsplit(" > ", 1)[-1]
    fixed = len("章节：") + len(suffix)
    available = max_chars - fixed - 1
    if available < len(last_section):
        suffix = ""
        available = max_chars - len("章节：") - 1
    if available <= 0:
        return last_section[-max_chars:]
    return f"章节：{last_section[-available:]}{suffix}"


def _chunk_text(
    document_id: UUID,
    document_name: str,
    content: str,
    item: dict,
    max_chars: int,
    overlap: int,
    *,
    prefix: str = "",
) -> list[DocumentChunk]:
    capacity = max_chars - len(prefix)
    if capacity <= 0:
        prefix = ""
        capacity = max_chars
    chunks: list[DocumentChunk] = []
    start = 0
    while start < len(content):
        end = _natural_end(content, start, capacity)
        chunks.append(_make_chunk(
            document_id, document_name, prefix + content[start:end], item,
            char_start=start, char_end=end,
        ))
        if end == len(content):
            break
        next_start = end - min(overlap, end - start)
        start = next_start if next_start > start else end
    return chunks


def _table_context(content: str, max_chars: int) -> tuple[str, list[str]]:
    lines = content.splitlines()
    if len(lines) < 3:
        return content[:max_chars - 1], []
    section = lines[0].removeprefix("章节：")
    metadata = f"\n{lines[1]}\n{lines[2]}"
    prefix = _prefix_with_section(section, metadata, max_chars)
    if prefix:
        return prefix, lines[3:]
    # 极小上限时优先保留末级标题，宁可省略无法容纳的表格元数据。
    return _prefix_with_section(section, "", max_chars), lines[3:]


def _chunk_table(
    document_id: UUID,
    document_name: str,
    content: str,
    item: dict,
    max_chars: int,
) -> list[DocumentChunk]:
    context, rows = _table_context(content, max_chars)
    chunks: list[DocumentChunk] = []
    pending: list[tuple[str, int, int]] = []
    source_offset = 0

    def emit(lines: list[tuple[str, int, int]]) -> None:
        body = "\n".join(line for line, _, _ in lines)
        part = context if not body else f"{context}\n{body}"
        chunks.append(_make_chunk(
            document_id, document_name, part, item,
            char_start=lines[0][1] if lines else 0,
            char_end=lines[-1][2] if lines else 0,
        ))

    for row in rows:
        row_start = source_offset
        row_end = row_start + len(row)
        source_offset = row_end + 1
        candidate = "\n".join(line for line, _, _ in pending + [(row, row_start, row_end)])
        if len(context) + 1 + len(candidate) <= max_chars:
            pending.append((row, row_start, row_end))
            continue
        if pending:
            emit(pending)
            pending = []
        if len(context) + 1 + len(row) <= max_chars:
            pending.append((row, row_start, row_end))
            continue
        label, separator, remainder = row.partition("：")
        if not separator:
            raise ValueError("表格行格式无效")
        source_cursor = row_start + len(label) + 1
        first = True
        while remainder:
            row_label = f"{label}{separator}" if first else f"{label}（续）："
            room = max_chars - len(context) - 1 - len(row_label)
            if room <= 0:
                row_label = ""
                room = max_chars - len(context) - 1
            if room <= 0:
                context = ""
                room = max_chars
            end = _natural_end(remainder, 0, room)
            part = remainder[:end]
            emit([(row_label + part, source_cursor, source_cursor + len(part))])
            source_cursor += len(part)
            remainder = remainder[end:]
            first = False
    if pending or not chunks:
        emit(pending)
    return chunks


def _structured_text_groups(items: list[dict]) -> list[list[dict]]:
    groups: list[list[dict]] = []
    for item in items:
        content = item.get("content")
        if not isinstance(content, str) or not content.strip():
            continue
        chunk_type = item.get("chunk_type")
        if chunk_type not in {"text", "table", "chart_ocr"}:
            raise ValueError("未知块类型")
        previous = groups[-1][-1] if groups else None
        orders_are_discontinuous = (
            isinstance(item.get("source_order"), int)
            and isinstance(previous.get("source_order"), int) if previous else False
        ) and item["source_order"] != previous["source_order"] + 1
        if (
            chunk_type == "text"
            and groups
            and groups[-1][0].get("chunk_type") == "text"
            and not orders_are_discontinuous
            and tuple(item.get("section_path", ())) == tuple(groups[-1][0].get("section_path", ()))
        ):
            groups[-1].append(item)
        else:
            groups.append([item])
    return groups


def chunk_extracted_content(
    document_id: UUID,
    document_name: str,
    items: list[dict],
    max_chars: int = MAX_CHARS,
    overlap: int = OVERLAP,
) -> list[DocumentChunk]:
    """按文档结构、字符上限和重叠规则将提取内容转换为可追溯块。"""
    if max_chars > MAX_CHARS:
        raise ValueError(f"max_chars 不得超过 {MAX_CHARS}")
    if overlap > OVERLAP:
        raise ValueError(f"overlap 不得超过 {OVERLAP}")
    if max_chars <= 0 or not 0 <= overlap < max_chars:
        raise ValueError("分块参数无效")

    chunks: list[DocumentChunk] = []
    for group in _structured_text_groups(items):
        item = group[0]
        chunk_type = item["chunk_type"]
        if chunk_type == "table":
            chunks.extend(_chunk_table(document_id, document_name, item["content"], item, max_chars))
            continue

        content = "\n\n".join(entry["content"] for entry in group)
        section = item.get("source_section")
        if not isinstance(section, str) or not section:
            section = _section_name(tuple(item.get("section_path", ())))
        if chunk_type == "chart_ocr":
            prefix = _prefix_with_section(section, f"\n图表 {item.get('source_figure', 1)}\n\n", max_chars)
        elif section != "正文":
            prefix = _prefix_with_section(section, "\n\n", max_chars)
        else:
            prefix = ""
        chunks.extend(_chunk_text(
            document_id, document_name, content, item, max_chars, overlap, prefix=prefix,
        ))
    return chunks


def _heading_level(style_name: str) -> int | None:
    if not style_name.startswith(("Heading", "标题")):
        return None
    match = re.search(r"(\d+)", style_name)
    return int(match.group(1)) if match else 0


def _single_line_cell_text(value: str) -> str:
    return " ".join(value.split())


def _format_table(table, section: str, table_index: int) -> str:
    header_index = next((index for index, row in enumerate(table.rows, start=1) if any(_single_line_cell_text(cell.text) for cell in row.cells)), None)
    if header_index is None:
        return ""
    header_cells = [_single_line_cell_text(cell.text) for cell in table.rows[header_index - 1].cells]
    headers = [value or f"第 {index} 列" for index, value in enumerate(header_cells, start=1)]
    lines = [f"章节：{section}", f"表格 {table_index}", f"表头：{' | '.join(headers)}"]
    data_row_index = 0
    for row in table.rows[header_index:]:
        values = [_single_line_cell_text(cell.text) for cell in row.cells]
        if not any(values):
            continue
        data_row_index += 1
        pairs = "；".join(f"{header}={value}" for header, value in zip(headers, values))
        lines.append(f"第 {data_row_index} 行：{pairs}")
    return "\n".join(lines)


def _extract_blips(document, element, output_dir: Path, *, base_item: dict, figure_index: int) -> tuple[list[dict], int]:
    images: list[dict] = []
    for blip in element.xpath(".//a:blip"):
        relation_id = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
        relation = document.part.rels.get(relation_id)
        if relation is None:
            continue
        figure_index += 1
        extension = Path(relation.target_ref).suffix or ".bin"
        image_name = f"figure-{figure_index}{extension.lower()}"
        image_bytes = relation.target_part.blob
        (output_dir / image_name).write_bytes(image_bytes)
        images.append(base_item | {
            "content": "", "chunk_type": "chart_ocr", "source_figure": figure_index,
            "image_path": image_name, "image_bytes": image_bytes,
        })
    return images, figure_index


def extract_docx(path: Path | str, extracted_dir: Path | str) -> list[dict]:
    """提取 DOCX 正文、表格和嵌入图片，并保留标题路径与 body 顺序。"""
    from docx import Document

    document = Document(path)
    output_dir = Path(extracted_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    items: list[dict] = []
    heading_path: list[str] = []
    table_index = 0
    figure_index = 0
    source_order = -1
    paragraphs_by_element = {paragraph._p: paragraph for paragraph in document.paragraphs}
    tables_by_element = {table._tbl: table for table in document.tables}

    def append(item: dict) -> None:
        nonlocal source_order
        source_order += 1
        items.append(item | {"source_order": source_order})

    for element in document.element.body.iterchildren():
        tag = element.tag.rsplit("}", 1)[-1]
        if tag == "p":
            paragraph = paragraphs_by_element[element]
            text = paragraph.text.strip()
            if text:
                level = _heading_level(paragraph.style.name)
                if level is not None:
                    heading_path = heading_path[:level - 1] if level else heading_path
                    heading_path.append(text)
                path_tuple = tuple(heading_path)
                append({
                    "content": text, "chunk_type": "text", "section_path": path_tuple,
                    "source_section": _section_name(path_tuple),
                })
            path_tuple = tuple(heading_path)
            base_item = {"section_path": path_tuple, "source_section": _section_name(path_tuple)}
            images, figure_index = _extract_blips(document, element, output_dir, base_item=base_item, figure_index=figure_index)
            for image in images:
                append(image)
        elif tag == "tbl":
            table = tables_by_element[element]
            table_index += 1
            path_tuple = tuple(heading_path)
            section = _section_name(path_tuple)
            content = _format_table(table, section, table_index)
            base_item = {"section_path": path_tuple, "source_section": section, "source_table": table_index}
            if content:
                append(base_item | {"content": content, "chunk_type": "table"})
            images, figure_index = _extract_blips(document, element, output_dir, base_item=base_item, figure_index=figure_index)
            for image in images:
                append(image)
    return items


def extract_pdf_with_pymupdf(path: Path | str, extracted_dir: Path | str) -> list[dict]:
    """使用 PyMuPDF 提取 PDF 文本和图片，作为 MinerU 的后备解析原语。"""
    import fitz

    output_dir = Path(extracted_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    items: list[dict] = []
    document = fitz.open(path)
    try:
        figure_index = 0
        for page_number, page in enumerate(document, start=1):
            text = page.get_text("text").strip()
            if text:
                items.append({"content": text, "chunk_type": "text", "source_page": page_number})
            for image in page.get_images(full=True):
                figure_index += 1
                image_data = document.extract_image(image[0])
                extension = image_data.get("ext", "bin")
                image_name = f"page-{page_number}-figure-{figure_index}.{extension}"
                (output_dir / image_name).write_bytes(image_data["image"])
                items.append({
                    "content": "PDF 图片", "chunk_type": "chart_ocr", "source_page": page_number,
                    "source_figure": figure_index, "image_path": image_name, "image_bytes": image_data["image"],
                })
    finally:
        document.close()
    return items
